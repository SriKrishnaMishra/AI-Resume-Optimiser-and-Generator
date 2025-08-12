from __future__ import annotations
import io
from fpdf import FPDF
from docx import Document
from docx.shared import Pt


def _latin1_safe(text: str) -> str:
    return text.encode("latin-1", "replace").decode("latin-1")


def to_pdf_bytes(text: str) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=11)

    max_w = pdf.w - pdf.l_margin - pdf.r_margin
    line_height = 6

    def break_long_word(word: str):
        parts = []
        buf = ""
        for ch in word:
            if pdf.get_string_width(buf + ch) <= max_w:
                buf += ch
            else:
                if buf:
                    parts.append(buf)
                    buf = ch
                else:
                    # Force progress even if a single character exceeds width (rare)
                    parts.append(ch)
                    buf = ""
        if buf:
            parts.append(buf)
        return parts

    def wrap_line(s: str):
        words = s.split(" ")
        lines = []
        current = ""
        while words:
            w = words[0]
            candidate = w if not current else current + " " + w
            if pdf.get_string_width(candidate) <= max_w:
                current = candidate
                words.pop(0)
            else:
                if current:
                    lines.append(current)
                    current = ""
                else:
                    # Word alone longer than max width: break it
                    parts = break_long_word(w)
                    lines.extend(parts[:-1])
                    last = parts[-1] if parts else ""
                    current = last
                    words.pop(0)
        if current:
            lines.append(current)
        return lines

    for raw in text.splitlines():
        safe = _latin1_safe(raw).replace("\t", "    ")
        if not safe.strip():
            pdf.ln(line_height // 2)
            continue
        for chunk in wrap_line(safe):
            pdf.set_x(pdf.l_margin)
            pdf.cell(0, line_height, chunk, ln=1)

    return pdf.output(dest="S").encode("latin-1", "ignore")


def to_docx_bytes(text: str) -> bytes:
    doc = Document()

    # Basic styling: detect simple section headers
    def is_header(line: str) -> bool:
        l = line.strip()
        if not l:
            return False
        return l.isupper() or l.endswith(":") or l in {"Professional Summary", "Key Skills", "Experience", "Education", "Projects"}

    for line in text.splitlines():
        if is_header(line):
            p = doc.add_heading(line.strip().rstrip(":"), level=2)
        else:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(10)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
